import streamlit as st
import sys
import os
import re
import docx
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import pandas as pd
import inspect
import importlib
import sqlite3  # <-- to catch IntegrityError
from typing import Optional

# ---- Feature flags ----
ENABLE_FLOATING_CHAT = False  # keep inline, not floating

# Make local packages importable when run from /ui
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ------------------------- Imports from your app -------------------------
from core.analyzer import (
    check_requirement_ambiguity,
    check_passive_voice,
    check_incompleteness,
)
try:
    from core.analyzer import check_singularity  # optional
except Exception:
    def check_singularity(_text: str):
        return []  # safe fallback

from core.scoring import calculate_clarity_score

try:
    from core.rule_engine import RuleEngine
except Exception:
    class RuleEngine:
        def __init__(self):
            pass

from llm import ai_suggestions as ai  # module import
ai = importlib.reload(ai)

get_ai_suggestion = getattr(ai, "get_ai_suggestion")
generate_requirement_from_need = getattr(ai, "generate_requirement_from_need")

get_chatbot_response = getattr(
    ai,
    "get_chatbot_response",
    lambda api_key, history: get_ai_suggestion(
        api_key,
        "\n".join(
            f"{m.get('role','user').upper()}: {(m.get('parts') or [m.get('content','')])[0]}"
            for m in history
        ) + "\nASSISTANT:"
    ),
)

decompose_requirement_with_ai = getattr(
    ai, "decompose_requirement_with_ai",
    lambda api_key, requirement_text: "Decomposition helper failed to load."
)

try:
    from llm.ai_suggestions import extract_requirements_with_ai
    HAS_AI_PARSER = True
except Exception:
    HAS_AI_PARSER = False
    def extract_requirements_with_ai(*args, **kwargs):
        return []

from db.database import init_db, add_project, get_all_projects  # type: ignore
from db import database as db  # type: ignore
db = importlib.reload(db)

# ---------------------- Theme from config.toml helpers ----------------------
def _hex_to_rgb(h: str):
    s = (h or "").lstrip("#")
    if len(s) == 3:
        s = "".join(c*2 for c in s)
    try:
        return int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16)
    except Exception:
        return (0, 0, 0)

def _is_light_bg(hex_color: str) -> bool:
    r, g, b = _hex_to_rgb(hex_color)
    # perceived luminance (ITU-R BT.601)
    return (0.299 * r + 0.587 * g + 0.114 * b) > 186

def _font_stack_from_config(name: str) -> str:
    name = (name or "").strip().lower()
    if name == "serif":
        return 'Georgia, "Times New Roman", serif'
    if name == "monospace":
        return 'SFMono-Regular, Menlo, Consolas, "Liberation Mono", monospace'
    # default "sans serif"
    return 'Inter, -apple-system, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif'

def _theme_from_config() -> dict:
    primary = st.get_option("theme.primaryColor") or "#4F46E5"
    bg = st.get_option("theme.backgroundColor") or "#111827"
    secondary = st.get_option("theme.secondaryBackgroundColor") or "#1F2937"
    text = st.get_option("theme.textColor") or "#F9FAFB"
    font_name = st.get_option("theme.font") or "sans serif"
    light = _is_light_bg(bg)
    border = "rgba(0,0,0,.12)" if light else "rgba(255,255,255,.14)"
    muted = "rgba(0,0,0,.55)" if light else "rgba(255,255,255,.7)"
    return {
        "primary": primary,
        "bg": bg,
        "secondary": secondary,
        "text": text,
        "border": border,
        "muted": muted,
        "font_stack": _font_stack_from_config(font_name),
    }

# --- API key validator (Google AI Studio keys commonly start with "AIza") ---
def is_valid_google_key(k: str) -> bool:
    k = (k or "").strip()
    if not k:
        return False
    if not k.startswith("AIza"):
        return False
    # heuristic: length + safe chars
    return 30 <= len(k) <= 120 and re.fullmatch(r"[A-Za-z0-9_\-]+", k) is not None

# ========================= Helpers for Analyzer =========================
def _read_docx_text_and_rows(uploaded_file):
    d = docx.Document(uploaded_file)
    parts, rows = [], []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in d.tables:
        for r in tbl.rows:
            row_cells = []
            for c in r.cells:
                cell_text = " ".join(p.text.strip() for p in c.paragraphs if p.text.strip())
                row_cells.append(cell_text)
                if cell_text:
                    parts.append(cell_text)
            rows.append(row_cells)
    return "\n".join(parts), rows

def _read_docx_text_and_rows_from_path(path: str):
    d = docx.Document(path)
    parts, rows = [], []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in d.tables:
        for r in tbl.rows:
            row_cells = []
            for c in r.cells:
                cell_text = " ".join(p.text.strip() for p in c.paragraphs if p.text.strip())
                row_cells.append(cell_text)
                if cell_text:
                    parts.append(cell_text)
            rows.append(row_cells)
    return "\n".join(parts), rows

def _extract_requirements_from_table_rows(table_rows):
    if not table_rows:
        return []
    def _norm(s): return (s or "").strip().lower()
    header_idx = None
    for i, row in enumerate(table_rows):
        cells = [_norm(c) for c in row]
        if not cells:
            continue
        if ("id" in cells[0] and any("requirement" in c for c in cells)):
            header_idx = i
            break
        if ("requirement" in cells[0] and any(c == "id" for c in cells)):
            header_idx = i
            break
    if header_idx is None:
        return []
    header = [_norm(c) for c in table_rows[header_idx]]
    id_col = req_col = None
    for idx, h in enumerate(header):
        if h == "id":
            id_col = idx
        if "requirement" in h:
            req_col = idx
    if id_col is None or req_col is None:
        return []
    id_pat = re.compile(r'^[A-Z][A-Z0-9-]*-\d+$')
    out = []
    for row in table_rows[header_idx + 1:]:
        if len(row) <= max(id_col, req_col):
            continue
        rid = (row[id_col] or "").strip()
        rtx = (row[req_col] or "").strip()
        if rid and rtx and id_pat.match(rid):
            out.append((rid, rtx))
    return out

def extract_requirements_from_string(content: str):
    requirements = []
    req_pattern = re.compile(r'^(([A-Z][A-Z0-9-]*-\d+)|(\d+\.))\s+(.*)$')
    for line in content.split('\n'):
        line = line.strip()
        m = req_pattern.match(line)
        if m:
            rid = m.group(1)
            text = m.group(4)
            requirements.append((rid, text))
    return requirements

def extract_requirements_from_file(uploaded_file):
    if uploaded_file.name.endswith('.txt'):
        content = uploaded_file.getvalue().decode("utf-8")
        table_rows = []
    elif uploaded_file.name.endswith('.docx'):
        content, table_rows = _read_docx_text_and_rows(uploaded_file)
    else:
        content, table_rows = "", []
    reqs = _extract_requirements_from_table_rows(table_rows)
    return reqs or extract_requirements_from_string(content)

def format_requirement_with_highlights(req_id, req_text, issues):
    highlighted_text = req_text

    # --- ambiguity highlighting ---
    if issues.get('ambiguous'):
        for token in issues['ambiguous']:
            word = token.split(":", 1)[1].strip() if ":" in token else token
            if not word:
                word = token
            highlighted_text = re.sub(
                r'\b' + re.escape(word) + r'\b',
                f'<span style="background-color:#FFFF00;color:black;padding:2px 4px;border-radius:3px;">{word}</span>',
                highlighted_text,
                flags=re.IGNORECASE
            )

    if issues.get('passive'):
        for phrase in issues['passive']:
            highlighted_text = re.sub(
                re.escape(phrase),
                f'<span style="background-color:#FFA500;padding:2px 4px;border-radius:3px;">{phrase}</span>',
                highlighted_text, flags=re.IGNORECASE
            )

    display_html = f"⚠️ <strong>{req_id}</strong> {highlighted_text}"
    explanations = []
    if issues.get('ambiguous'):
        explanations.append(f"<i>- Ambiguity: Found weak words: <b>{', '.join(issues['ambiguous'])}</b></i>")
    if issues.get('passive'):
        explanations.append(f"<i>- Passive Voice: Found phrase: <b>'{', '.join(issues['passive'])}'</b>. Consider active voice.</i>")
    if issues.get('incomplete'):
        explanations.append("<i>- Incompleteness: Requirement appears to be a fragment.</i>")
    if issues.get('singularity'):
        explanations.append(f"<i>- Singularity: Multiple actions indicated: <b>{', '.join(issues['singularity'])}</b></i>")
    if explanations:
        display_html += "<br>" + "<br>".join(explanations)

    return (
        f'<div style="background-color:#FFF3CD;color:#856404;padding:10px;'
        f'border-radius:5px;margin-bottom:10px;">{display_html}</div>'
    )

def safe_call_ambiguity(text: str, engine: Optional['RuleEngine']):
    """
    Prefer the JSON-driven RuleEngine.check_ambiguity() so new rules apply.
    Record which path was used for quick diagnostics.
    """
    # 1) Try engine path
    try:
        if engine and hasattr(engine, "check_ambiguity"):
            out = engine.check_ambiguity(text or "") or []
            st.session_state["dbg_ambiguity_path"] = "engine"
            st.session_state["dbg_last_amb"] = out
            return out
    except Exception as e:
        st.session_state["dbg_ambiguity_error"] = f"engine: {e}"

    # 2) Legacy fallback
    try:
        out = check_requirement_ambiguity(text, engine)
        st.session_state["dbg_ambiguity_path"] = "legacy"
        st.session_state["dbg_last_amb"] = out
        return out
    except TypeError:
        out = check_requirement_ambiguity(text)
        st.session_state["dbg_ambiguity_path"] = "legacy(no-engine-param)"
        st.session_state["dbg_last_amb"] = out
        return out
    except Exception as e:
        st.session_state["dbg_ambiguity_error"] = f"legacy: {e}"
        return []

def safe_clarity_score(total_reqs: int, results: list[dict], issue_counts=None, engine: Optional['RuleEngine']=None):
    try:
        sig = inspect.signature(calculate_clarity_score)
        if len(sig.parameters) >= 3:
            return calculate_clarity_score(total_reqs, issue_counts or {}, engine)
        else:
            flagged_reqs = sum(1 for r in results if r['ambiguous'] or r['passive'] or r['incomplete'])
            return calculate_clarity_score(total_reqs, flagged_reqs)
    except Exception:
        flagged_reqs = sum(1 for r in results if r['ambiguous'] or r['passive'] or r['incomplete'])
        clear_reqs = max(0, total_reqs - flagged_reqs)
        return int((clear_reqs / total_reqs) if total_reqs else 1 * 100)

def _open_db_conn():
    if hasattr(db, "get_connection"):
        try:
            conn = db.get_connection()
            conn.execute("PRAGMA foreign_keys = ON;")
            return conn, False
        except Exception:
            pass
    for attr in ("DB_PATH", "DB_FILE", "DB_NAME", "DATABASE_PATH"):
        path = getattr(db, attr, None)
        if isinstance(path, str) and path:
            conn = sqlite3.connect(path)
            conn.execute("PRAGMA foreign_keys = ON;")
            return conn, True
    base_dir = os.path.dirname(db.__file__)
    for name in ("reqcheck.db", "database.db", "app.db"):
        candidate = os.path.join(base_dir, name)
        if os.path.exists(candidate):
            conn = sqlite3.connect(candidate)
            conn.execute("PRAGMA foreign_keys = ON;")
            return conn, True
    raise RuntimeError("Could not locate the SQLite DB file. Expose DB_PATH in db.database or provide db.get_connection().")

def _sanitize_filename(name: str) -> str:
    return re.sub(r'[^A-Za-z0-9._-]+', '_', name)

def _save_uploaded_file_for_doc(project_id: int, doc_id: int, display_name: str, uploaded_file) -> str:
    base_dir = os.path.join("data", "projects", str(project_id), "documents")
    os.makedirs(base_dir, exist_ok=True)
    safe_name = _sanitize_filename(display_name)
    out_path = os.path.join(base_dir, f"{doc_id}_{safe_name}")
    try:
        uploaded_file.seek(0)
        with open(out_path, "wb") as f:
            f.write(uploaded_file.read())
    except Exception:
        with open(out_path, "wb") as f:
            f.write(uploaded_file.getvalue())
    return out_path

# =============================== UI Setup ===============================
st.set_page_config(
    page_title="ReqCheck Workspace",
    page_icon="https://github.com/vin-2020/Requirements-Clarity-Checker/blob/main/Logo.png?raw=true",
    layout="wide"
)

# -- small utility styles used by analyzer outputs (unchanged app logic) --
st.markdown("""
<style>
    .req-container { padding:10px;border-radius:5px;margin-bottom:10px;border:1px solid #ddd; }
    .flagged { background:#FFF3CD;color:#856404;border-color:#FFEEBA; }
    .clear { background:#D4EDDA;color:#155724;border-color:#C3E6CB; }
    .highlight-ambiguity { background:#FFFF00;color:black;padding:2px 4px;border-radius:3px; }
    .highlight-passive { background:#FFA500;padding:2px 4px;border-radius:3px; }
    .explanation { font-size:0.9em;font-style:italic;color:#6c757d;margin-top:5px; }
</style>
""", unsafe_allow_html=True)

# ---- Theme backbone (colors now come from config.toml) ----
def apply_theme_vars():
    t = _theme_from_config()
    st.markdown(f"""
    <style>
      :root {{
        --rc-primary: {t['primary']};
        --rc-bg: {t['bg']};
        --rc-bg-2: {t['secondary']};
        --rc-text: {t['text']};
        --rc-border: {t['border']};
        --rc-muted: {t['muted']};
      }}
      html, body, [data-testid="stAppViewContainer"] {{
        background: var(--rc-bg) !important;
        color: var(--rc-text) !important;
        font-family: {t['font_stack']} !important;
      }}
      a, a:visited {{ color: var(--rc-primary) !important; }}
    </style>
    """, unsafe_allow_html=True)

apply_theme_vars()  # Apply once on load

# ---- Single API Key Card (status + input + link) ----
if "api_key" not in st.session_state:
    st.session_state.api_key = ""

typed_key = st.session_state.get("api_key_global", st.session_state.get("api_key", ""))
is_connected = is_valid_google_key(typed_key)

# Card styling (scoped)
st.markdown("""
<style>
.api-card {
  background: var(--rc-bg-2);
  border: 1px solid var(--rc-border);
  border-radius: 12px;
  padding: 14px;
  margin: 6px 0 14px;
}
.api-row { display: grid; grid-template-columns: auto 1fr auto; gap: 12px; align-items: center; }
.status-pill {
  display:inline-flex; align-items:center; gap:8px; padding:6px 12px;
  border-radius: 999px; font-weight:700; font-size:13px; border:1px solid var(--rc-border);
}
.status-dot { width:8px; height:8px; border-radius:50%; display:inline-block; }
.status-ok  { background:#DCFCE7; color:#065F46; border-color:rgba(16,185,129,.35); }
.status-ok .status-dot { background:#10B981; }
.status-bad { background:#FEF3C7; color:#92400E; border-color:rgba(234,179,8,.35); }
.status-bad .status-dot { background:#EAB308; }
.api-help { margin-top: 8px; opacity:.85; }
</style>
""", unsafe_allow_html=True)

with st.container(border=False):
    st.markdown('<div class="api-card">', unsafe_allow_html=True)
    st.markdown('<div class="api-row">', unsafe_allow_html=True)
    # Left: status pill
    pill_class = "status-ok" if is_connected else "status-bad"
    pill_text  = "Connected" if is_connected else "Disconnected"
    st.markdown(
        f'<div class="status-pill {pill_class}"><span class="status-dot"></span>{pill_text}</div>',
        unsafe_allow_html=True
    )
    # Middle: API key input
    st.text_input(
        "API key",
        key="api_key_global",
        value=typed_key,
        type="password",
        label_visibility="collapsed",
        help="Your key stays in this session only."
    )
    st.session_state.api_key = st.session_state.api_key_global
    # Right: link button
    st.link_button("Get key — Google AI Studio", "https://aistudio.google.com/", use_container_width=False)
    st.markdown('</div>', unsafe_allow_html=True)  # /.api-row
    st.markdown(' <div class="api-help">Tip: paste a key starting with <code>AIza</code> to connect.</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)  # /.api-card

# Track selected project globally
if 'selected_project' not in st.session_state:
    st.session_state.selected_project = None

# One RuleEngine instance (real or stub)
rule_engine = RuleEngine()

# ======================= Layout: main + right panel =======================
main_col, right_col = st.columns([4, 1], gap="large")

# ----------------------------- Right Panel (Projects) -----------------------------
with right_col:
    st.subheader("🗂️ Projects")

    if st.session_state.selected_project is not None:
        _pid, _pname = st.session_state.selected_project
        st.caption(f"Current: **{_pname}**")
        if st.button("Clear selection", key="btn_clear_proj_right"):
            st.session_state.selected_project = None
            st.rerun()

    projects = get_all_projects()
    names = [p[1] for p in projects] if projects else []
    if names:
        sel_name = st.selectbox("Open project:", names, key="proj_select_right")

        if "confirm_delete" not in st.session_state:
            st.session_state.confirm_delete = False
        if "delete_project_id" not in st.session_state:
            st.session_state.delete_project_id = None
        if "delete_project_name" not in st.session_state:
            st.session_state.delete_project_name = None

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Load", key="btn_load_proj_right"):
                for p in projects:
                    if p[1] == sel_name:
                        st.session_state.selected_project = p
                        st.success(f"Loaded: {sel_name}")
                        st.rerun()
        with col2:
            if st.button("Delete", key="btn_delete_proj_right"):
                for p in projects:
                    if p[1] == sel_name:
                        st.session_state.delete_project_id = p[0]
                        st.session_state.delete_project_name = sel_name
                        st.session_state.confirm_delete = True

        if st.session_state.confirm_delete:
            st.warning(
                f"You're about to delete '{st.session_state.delete_project_name}'. "
                "This cannot be undone."
            )
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Confirm Delete", key="btn_confirm_delete_proj_right"):
                    pid_to_delete = st.session_state.delete_project_id

                    def _get_tables(conn):
                        cur = conn.cursor()
                        cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
                        return [r[0] for r in cur.fetchall()]

                    def _fk_refs(conn, ref_table_name):
                        refs = []
                        for t in _get_tables(conn):
                            try:
                                cur = conn.cursor()
                                cur.execute(f"PRAGMA foreign_key_list('{t}')")
                                for (id_, seq, table, from_col, to_col, on_update, on_delete, match) in cur.fetchall():
                                    if table == ref_table_name:
                                        refs.append((t, from_col))
                            except Exception:
                                pass
                        return refs

                    try:
                        conn, _should_close = _open_db_conn()
                        cur = conn.cursor()
                        cur.execute("PRAGMA foreign_keys = ON;")

                        cur.execute("SELECT id FROM documents WHERE project_id = ?", (pid_to_delete,))
                        doc_ids = [r[0] for r in cur.fetchall()]

                        if doc_ids:
                            refs_to_documents = _fk_refs(conn, "documents")
                            for (tbl, col) in refs_to_documents:
                                qmarks = ",".join("?" for _ in doc_ids)
                                cur.execute(f"DELETE FROM {tbl} WHERE {col} IN ({qmarks})", doc_ids)

                            qmarks = ",".join("?" for _ in doc_ids)
                            try:
                                cur.execute(f"DELETE FROM requirements WHERE document_id IN ({qmarks})", doc_ids)
                            except Exception:
                                pass
                            cur.execute(f"DELETE FROM documents WHERE id IN ({qmarks})", doc_ids)

                        refs_to_projects = _fk_refs(conn, "projects")
                        for (tbl, col) in refs_to_projects:
                            cur.execute(f"DELETE FROM {tbl} WHERE {col} = ?", (pid_to_delete,))

                        cur.execute("DELETE FROM projects WHERE id = ?", (pid_to_delete,))
                        conn.commit()
                        if _should_close:
                            conn.close()

                        st.success("Project deleted.")

                    except sqlite3.IntegrityError as e:
                        try:
                            cur.execute("PRAGMA foreign_key_check;")
                            problems = cur.fetchall()
                        except Exception:
                            problems = []
                        if problems:
                            st.error(f"Delete failed due to FK constraints. Offending rows: {problems}")
                        else:
                            st.error(f"Delete failed due to foreign key constraints: {e}")
                    except Exception as e:
                        st.error(f"Delete failed: {e}")

                    st.session_state.confirm_delete = False
                    st.session_state.delete_project_id = None
                    st.session_state.delete_project_name = None
                    st.rerun()

            with c2:
                if st.button("Cancel", key="btn_cancel_delete_proj_right"):
                    st.session_state.confirm_delete = False
                    st.session_state.delete_project_id = None
                    st.session_state.delete_project_name = None

    else:
        st.caption("No projects yet.")

    st.text_input("New project name:", key="new_proj_name_right")
    if st.button("Create", key="btn_create_proj_right"):
        new_name = st.session_state.get("new_proj_name_right", "").strip()
        if new_name:
            feedback = add_project(new_name)
            st.success(feedback)
            st.rerun()
        else:
            st.error("Please enter a project name.")

    # ---------------- Quick Chat: inline under Projects (expander) ----------------
    st.markdown("---")
    with st.expander("🤖 Quick Chat", expanded=True):
        try:
            import ui.widgets.quick_chat as _quick_chat  # import the module
            _quick_chat = importlib.reload(_quick_chat)  # hot reload to avoid stale function
            if hasattr(_quick_chat, "render_inline_quick_chat"):
                _quick_chat.render_inline_quick_chat(st, {
                    "get_chatbot_response": get_chatbot_response
                })
            else:
                st.error("render_inline_quick_chat() not found in ui/widgets/quick_chat.py")
        except Exception as e:
            st.error(f"Quick Chat failed to load: {e}")

# --- Website-like global navbar (logo left, tabs right) ---
st.markdown("""
<style>
/* === NAVBAR (fixed top) === */
.rc-navbar {
  position: fixed;
  top: 0;
  left: 0;
  right: 0;
  z-index: 99999; /* keep it above everything */
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 14px;
  padding: 24px 34px;
  background: var(--rc-bg);
  border-bottom: 1px solid var(--rc-border);
  backdrop-filter: blur(8px);
  box-shadow: 0 2px 8px rgba(0,0,0,0.08);
}

/* Override Streamlit’s default header offset */
[data-testid="stHeader"] { z-index: 0 !important; }
 
/* Push app content below fixed navbar */
body, [data-testid="stAppViewContainer"] > div:first-child {
  padding-top: 90px !important; /* was 60px */
}
 
 /* Brand (left) */
 .rc-brand { display:flex; align-items:center; gap:14px; }
 .rc-logo { height: 40px; width:auto; border-radius:12px; }
 .rc-name {
   font-family: "Poppins", Inter, -apple-system, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
   font-weight: 800; font-size: 22px;
   background: linear-gradient(90deg, var(--rc-primary), #22D3EE);
   -webkit-background-clip: text; -webkit-text-fill-color: transparent;
   letter-spacing: .3px;
 }
 /* Menu (right) */
 .rc-menu { display:flex; align-items:center; gap: 18px; flex-wrap: wrap; }
 .rc-link {
   display:inline-flex; align-items:center; gap:8px;
   font-weight: 600; font-size: 15px; color: var(--rc-text);
   text-decoration: none; cursor: pointer; position: relative; padding: 6px 2px;
 }
 .rc-link:hover { color: var(--rc-primary); }
 /* Active underline */
 .rc-link.active::after {
   content:""; position:absolute; left:0; right:0; bottom:-8px;
   height: 3px; background: var(--rc-primary); border-radius: 4px;
 }
 @media (max-width: 800px) { .rc-menu { gap: 12px; } }
</style>
 
<div class="rc-navbar">
  <div class="rc-brand">
    <img class="rc-logo" src="https://github.com/vin-2020/Requirements-Clarity-Checker/blob/main/ReqCheck_Logo.png?raw=true" alt="ReqCheck"/>
    <div class="rc-name">ReqCheck</div>
  </div>
  <nav class="rc-menu" id="rc-menu">
    <a class="rc-link" data-tab="0">🏠 Home</a>
    <a class="rc-link" data-tab="1">📄 Analyzer</a>
    <a class="rc-link" data-tab="2">💡 Need→Req</a>
    <a class="rc-link" data-tab="3">💬 Chatbot</a>
  </nav>
</div>
""", unsafe_allow_html=True)

with main_col:
    tab_home, tab_analyze, tab_need, tab_chat = st.tabs([
        "🏠 Home",
        "📄 Document Analyzer",
        "💡 Need-to-Requirement Helper",
        "💬 Requirements Chatbot",
    ])

    # Hide Streamlit's default tab strip — we will drive tabs via the navbar
    st.markdown("""
    <style> div[role="tablist"] { display: none !important; } </style>
    """, unsafe_allow_html=True)

    # Wire navbar clicks to hidden tabs + keep active highlight in sync
    from streamlit.components.v1 import html as _html
    _html("""
    <script>
    (function(){
      const P = window.parent || window;
      const doc = P.document;
      function getTabs(){ return Array.from(doc.querySelectorAll('button[role="tab"]')); }
      function setActive(idx){
        const links = doc.querySelectorAll('.rc-menu .rc-link');
        links.forEach(el => el.classList.remove('active'));
        if (links[idx]) links[idx].classList.add('active');
      }
      // Click on navbar -> click hidden tab
      doc.querySelectorAll('.rc-menu .rc-link').forEach(link=>{
        link.addEventListener('click', ()=>{
          const idx = Number(link.getAttribute('data-tab')) || 0;
          const tabs = getTabs();
          if (tabs[idx] && typeof tabs[idx].click === 'function') tabs[idx].click();
          setActive(idx);
        }, {passive:true});
      });
      // Init active state
      setActive(0);
      // Observe tab selection changes (keyboard, code, etc.)
      const tablist = doc.querySelector('div[role="tablist"]');
      if (tablist && 'MutationObserver' in P){
        new MutationObserver(()=>{
          const tabs = getTabs();
          const activeIdx = tabs.findIndex(b => b.getAttribute('aria-selected') === 'true');
          if (activeIdx >= 0) setActive(activeIdx);
        }).observe(tablist, { attributes:true, subtree:true, attributeFilter:['aria-selected','class'] });
      }
    })();
    </script>
    """, height=0)

# ------------------------------ Main Tabs (LEAN loader) ------------------------------
home_tab = analyzer_tab = needs_tab = chat_tab = None

try:
    import ui.tabs.home_tab as _home_tab
    home_tab = importlib.reload(_home_tab)
except Exception as e:
    st.error(f"Home tab failed to import: {e}")

try:
    import ui.tabs.analyzer_tab as _analyzer_tab
    analyzer_tab = importlib.reload(_analyzer_tab)
except Exception as e:
    st.error(f"Analyzer tab failed to import: {e}")

try:
    import ui.tabs.needs_tab as _needs_tab
    needs_tab = importlib.reload(_needs_tab)
except Exception as e:
    st.error(f"Needs tab failed to import: {e}")

try:
    import ui.tabs.chat_tab as _chat_tab
    chat_tab = importlib.reload(_chat_tab)
except Exception as e:
    st.error(f"Chat tab failed to import: {e}")

CTX = {
    "HAS_AI_PARSER": HAS_AI_PARSER,
    "get_ai_suggestion": get_ai_suggestion,
    "get_chatbot_response": get_chatbot_response,
    "decompose_requirement_with_ai": decompose_requirement_with_ai,
    "decompose_need_into_requirements": ai.decompose_need_into_requirements,
    "run_freeform": ai.run_freeform,
    "extract_requirements_with_ai": extract_requirements_with_ai,
    "_read_docx_text_and_rows": _read_docx_text_and_rows,
    "_read_docx_text_and_rows_from_path": _read_docx_text_and_rows_from_path,
    "_extract_requirements_from_table_rows": _extract_requirements_from_table_rows,
    "extract_requirements_from_string": extract_requirements_from_string,
    "extract_requirements_from_file": extract_requirements_from_file,
    "format_requirement_with_highlights": format_requirement_with_highlights,
    "safe_call_ambiguity": safe_call_ambiguity,
    "check_passive_voice": check_passive_voice,
    "check_incompleteness": check_incompleteness,
    "check_singularity": check_singularity,
    "safe_clarity_score": safe_clarity_score,
    "_save_uploaded_file_for_doc": _save_uploaded_file_for_doc,
    "_sanitize_filename": _sanitize_filename,
}

with tab_home:
    if home_tab:
        home_tab.render(st, db, rule_engine, CTX)
    else:
        st.error("Home tab not available.")

with tab_analyze:
    if analyzer_tab:
        analyzer_tab.render(st, db, rule_engine, CTX)
    else:
        st.error("Analyzer tab not available.")

with tab_need:
    if needs_tab:
        needs_tab.render(st, db, rule_engine, CTX)
    else:
        st.error("Need-to-Requirement tab not available.")

with tab_chat:
    if chat_tab:
        chat_tab.render(st, db, rule_engine, CTX)
    else:
        st.error("Chat tab not available.")

# ------------------------------ Footer (useful links at bottom) ------------------------------
st.markdown("""
<style>
.footer {
  margin-top: 40px;
  padding: 16px 0;
  border-top: 1px solid var(--rc-border);
  text-align: center;
  font-size: 14px;
  opacity: 0.95;
}
.footer a {
  color: var(--rc-primary) !important;
  font-weight: 600;
  text-decoration: none;
  margin: 0 10px;
}
.footer a:hover { text-decoration: underline; }
</style>

<div class="footer">
  <a href="https://aistudio.google.com/" target="_blank">Google AI Studio</a> •
  <a href="https://github.com/vin-2020/Requirements-Clarity-Checker" target="_blank">GitHub</a> •
  <a href="https://www.incose.org/products-and-publications/se-handbook" target="_blank">INCOSE Handbook</a> •
  <a href="mailto:reqcheck.dev@gmail.com">Contact</a>
</div>
""", unsafe_allow_html=True)
